"""
Dokumentum betöltők PDF, DOCX, TXT és MD fájlokhoz.
A cél: strukturált szöveget kinyerni, ami alkalmas LLM-alapú chunkolásra.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class Document:
    """Betöltött dokumentum reprezentációja."""

    content: str
    file_path: str
    file_type: str
    metadata: dict = field(default_factory=dict)


def load_pdf(file_path: Path) -> Document:
    """
    PDF fájl betöltése pymupdf4llm-mel.
    Markdown formátumú kimenetet ad, megőrzi a fejléceket, táblázatokat.
    """
    import pymupdf4llm

    md_text = pymupdf4llm.to_markdown(str(file_path))

    return Document(
        content=md_text,
        file_path=str(file_path),
        file_type="pdf",
        metadata={
            "source": str(file_path),
            "filename": file_path.name,
        },
    )


def load_docx(file_path: Path) -> Document:
    """
    DOCX fájl betöltése python-docx-szel.
    Paragrafusonként, stílusinformációkkal.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))

    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Heading stílusok Markdown fejlécekké alakítása
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading 1"):
            parts.append(f"# {text}")
        elif style_name.startswith("Heading 2"):
            parts.append(f"## {text}")
        elif style_name.startswith("Heading 3"):
            parts.append(f"### {text}")
        elif style_name.startswith("Heading"):
            parts.append(f"#### {text}")
        else:
            parts.append(text)

    # Táblázatok feldolgozása
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")

        if rows:
            # Markdown táblázat fejléc elválasztó
            header_sep = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            rows.insert(1, header_sep)
            parts.append("\n".join(rows))

    content = "\n\n".join(parts)

    return Document(
        content=content,
        file_path=str(file_path),
        file_type="docx",
        metadata={
            "source": str(file_path),
            "filename": file_path.name,
        },
    )


def load_txt(file_path: Path) -> Document:
    """TXT fájl betöltése UTF-8 kódolással."""
    content = file_path.read_text(encoding="utf-8")

    return Document(
        content=content,
        file_path=str(file_path),
        file_type="txt",
        metadata={
            "source": str(file_path),
            "filename": file_path.name,
        },
    )


def _clean_markdown(text: str) -> str:
    """
    Markdown szöveg tisztítása RAG indexeléshez.
    Eltávolítja a linkeket, képeket, HTML-t, lábjegyzeteket és egyéb
    metaadat szekciókat, amelyek zavarják a szemantikus chunkolást.
    """
    # --- HTML elemek eltávolítása ---
    # Video, table, és egyéb HTML tag-ek
    text = re.sub(r"<[^>]+>", "", text)

    # --- Képek eltávolítása ---
    # ![alt text](url) formátumú képhivatkozások
    # A (?:\\([^)]*\\)|[^)])* kezeli az escaped zárójeleket a Wikipedia URL-ekben
    text = re.sub(r"!\[[^\]]*\]\((?:\\[()]+|[^)])*\)", "", text)

    # --- Linkek szöveggé alakítása ---
    # [megjelenő szöveg](url "title") → megjelenő szöveg
    # Kezeli a Wikipedia-stílusú URL-eket escaped zárójelekkel és link title-ökkel
    text = re.sub(r"\[([^\]]+)\]\((?:\\[()]+|[^)])*\)", r"\1", text)

    # --- Maradék link title töredékek eltávolítása ---
    # Pl. "A Gyűrűk Ura (filmsorozat)") → üres
    text = re.sub(r'\s*"[^"]*"\)', "", text)

    # --- Lábjegyzet-hivatkozások eltávolítása ---
    # [^1], [^2] stb. inline hivatkozások
    text = re.sub(r"\[\^\d+\]", "", text)

    # --- Lábjegyzet-definíciók eltávolítása ---
    # Sorok amelyek [^N]: -tal kezdődnek (és a folytatásuk)
    text = re.sub(r"^\[\^\d+\]:.*$", "", text, flags=re.MULTILINE)

    # --- Irreleváns szekciók eltávolítása ---
    # Szekciók mint "További információk", "Jegyzetek", "Források",
    # "Kapcsolódó szócikkek", "Fordítás", "Megjegyzések"
    # Ezek a szekciók a szekció fejlécétől a következő azonos vagy magasabb
    # szintű fejlécig vagy a fájl végéig tartanak.
    irrelevant_sections = [
        r"További információk",
        r"Jegyzetek",
        r"Források",
        r"Kapcsolódó szócikkek",
        r"Fordítás",
        r"Megjegyzések",
        r"Külső hivatkozások",
        r"Lásd még",
        r"Hivatkozások",
        r"Bibliográfia",
        r"Irodalom",
        r"References",
        r"External links",
        r"See also",
        r"Further reading",
        r"Notes",
    ]
    for section_name in irrelevant_sections:
        # Megkeresi a szekció fejlécét (bármilyen szintű ##)
        # és eltávolítja a következő azonos/magasabb szintű fejlécig
        pattern = (
            r"^(#{1,6})\s+"
            + section_name
            + r"\s*$"
            + r"(.*?)"
            + r"(?=^\1(?!#)\s|\Z)"
        )
        text = re.sub(pattern, "", text, flags=re.MULTILINE | re.DOTALL)

    # --- ISBN hivatkozások egyszerűsítése ---
    # [ISBN 0-618-34399-7](url) → ISBN 0-618-34399-7
    text = re.sub(
        r"\[ISBN\s+([\d-]+)\]\([^)]*\)", r"ISBN \1", text
    )

    # --- Üres sorok normalizálása ---
    # Több egymást követő üres sor → max 2 üres sor
    text = re.sub(r"\n{4,}", "\n\n\n", text)

    # Sorok eleji/végi felesleges szóközök
    lines = [line.rstrip() for line in text.split("\n")]
    text = "\n".join(lines)

    # Kezdő és záró üres sorok
    text = text.strip()

    return text


def load_md(file_path: Path) -> Document:
    """
    Markdown fájl betöltése és tisztítása RAG indexeléshez.
    Eltávolítja a linkeket, képeket, HTML-t, lábjegyzeteket
    és a metaadat szekciókat (Jegyzetek, Források, stb.).
    """
    raw_content = file_path.read_text(encoding="utf-8")
    content = _clean_markdown(raw_content)

    return Document(
        content=content,
        file_path=str(file_path),
        file_type="md",
        metadata={
            "source": str(file_path),
            "filename": file_path.name,
        },
    )


def load_document(file_path: Path) -> Document:
    """
    Automatikus fájlformátum felismerés és betöltés.
    Támogatott: .pdf, .docx, .txt, .md
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".txt": load_txt,
        ".md": load_md,
    }

    loader = loaders.get(suffix)
    if loader is None:
        raise ValueError(
            f"Nem támogatott fájlformátum: {suffix} ({file_path.name}). "
            f"Támogatott: {', '.join(loaders.keys())}"
        )

    return loader(file_path)
